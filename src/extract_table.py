from docx import Document
import csv
import sys

def extract_tables_to_csv(docx_path, csv_path):
    try:
        print(f"Opening document: {docx_path}")
        doc = Document(docx_path)
        
        print(f"Number of tables found: {len(doc.tables)}")
        
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            for i, table in enumerate(doc.tables):
                print(f"Processing table {i+1}")
                # Write headers
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                writer.writerow(headers)
                print(f"Headers: {headers}")
                
                # Write data rows
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    writer.writerow(row_data)
                    print(f"Row: {row_data}")
                    
        print(f"CSV file saved to: {csv_path}")
                    
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        raise

if __name__ == '__main__':
    extract_tables_to_csv('docs/Master_File_with_Pointers.docx', 'docs/master_file_table.csv')
